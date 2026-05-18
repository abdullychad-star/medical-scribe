import streamlit as st
from groq import Groq
import os
import tempfile
import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

st.set_page_config(page_title="Medical Scribe", page_icon="🩺", layout="wide")

client = Groq(api_key=st.secrets["GROQ_API_KEY"])

st.markdown("""
    <h1 style='font-size:28px; font-weight:600;'>🩺 Autonomous Medical Scribe</h1>
    <p style='color:gray;'>Whisper + Groq AI &nbsp;·&nbsp; For demonstration purposes only</p>
    <hr>
""", unsafe_allow_html=True)

st.subheader("Step 1 — Upload or Record Audio")
audio_file = st.file_uploader("Upload audio file", type=["mp3", "wav", "m4a"])
st.caption("Recording directly in the browser coming soon")

if audio_file:
    st.audio(audio_file)

    if st.button("▶ Generate Notes", type="primary"):

        with st.spinner("Transcribing audio..."):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
                tmp.write(audio_file.read())
                tmp_path = tmp.name

            with open(tmp_path, "rb") as f:
                transcription = client.audio.transcriptions.create(
                    file=f,
                    model="whisper-large-v3",
                    response_format="verbose_json",
                    language="en"
                )

            transcript = transcription.text
            transcript = transcript.split("That's the end of the consultation")[0]
            transcript = transcript.replace(" Music", "").strip()

        st.success("Transcription complete!")
        with st.expander("📄 View Transcript"):
            st.write(transcript)

        with st.spinner("Generating SOAP note..."):
            soap_prompt = f"""You are a senior physician writing a SOAP note for a colleague.

Rules:
- Use proper medical terminology
- In Subjective, follow OLDCARTS format (Onset, Location, Duration, Character, Alleviating/Aggravating factors, Radiation, Timing, Severity)
- In Assessment, rank diagnoses from MOST to LEAST likely with brief reasoning for each
- In Assessment, only reference symptoms explicitly stated in the transcript. Double-check each sentence before including it.
- In Plan, be specific — include actual medication names/doses, specific tests, and exact follow-up timeframe
- In Plan, if pericarditis is the leading diagnosis, always include: Ibuprofen 600mg TID x 2 weeks with food, Colchicine 0.5mg BID x 3 months, restrict strenuous activity, and follow up ECG in 1 week.
- Always include an "Allergies" line in Subjective after PMH. If no allergies were mentioned, write "Allergies: Not reported - verify before prescribing"
- If the patient was vague or unclear about something, note it explicitly with [Patient unclear - needs clarification] rather than guessing or omitting it
- If the doctor asked a question and the patient didn't give a clear answer, flag it in the relevant section
- If something wasn't mentioned, write "Not reported"
- Never invent information not in the transcript
- CRITICAL: Never include clinical findings or symptoms that were not explicitly stated in the transcript. If you are unsure, write "Not reported".
- CRITICAL: Before writing each sentence in the Assessment, ask yourself "did the patient or doctor explicitly say this in the transcript?" If the answer is no, do not include it.
- CRITICAL: Do not include any information in the Assessment that was not explicitly stated in the transcript. Hallucinated details in a medical note are dangerous and unacceptable.
- CRITICAL: Classic textbook symptoms commonly associated with a diagnosis must NOT be included unless the patient explicitly stated them in the transcript.

Format EXACTLY like this:

SUBJECTIVE:
Chief Complaint: [one sentence]
HPI: [OLDCARTS structured paragraph]
PMH: [past medical history or "Not reported"]
Allergies: [or "Not reported - verify before prescribing"]
Family History: [or "Not reported"]
Social History: [smoking, drinking, drugs, occupation]
ROS: [relevant systems reviewed]

OBJECTIVE:
Vitals: [or "Not reported"]
Physical Exam: [or "Not reported"]

ASSESSMENT:
1. [Most likely diagnosis] - [one sentence reasoning]
2. [Second diagnosis] - [one sentence reasoning]
3. [Third diagnosis] - [one sentence reasoning]

PLAN:
1. Diagnostics: [specific tests ordered]
2. Treatment: [specific medications with doses if applicable]
3. Referrals: [or "None at this time"]
4. Patient Education: [specific instructions]
5. Follow-up: [specific timeframe]

FINAL WARNING: Do not add any clinical details not explicitly stated below. Invented details in a medical note can harm patients.

Transcript:
{transcript}"""

            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": soap_prompt}]
            )
            soap_note = response.choices[0].message.content

        st.subheader("📋 SOAP Note")
        sections = {"SUBJECTIVE": "", "OBJECTIVE": "", "ASSESSMENT": "", "PLAN": ""}
        current = None
        for line in soap_note.split("\n"):
            for section in sections:
                if line.strip().startswith(section):
                    current = section
                    break
            if current:
                sections[current] += line + "\n"

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Subjective**")
            st.text_area("", sections["SUBJECTIVE"], height=200, key="s", label_visibility="collapsed")
            st.markdown("**Objective**")
            st.text_area("", sections["OBJECTIVE"], height=100, key="o", label_visibility="collapsed")
        with col2:
            st.markdown("**Assessment**")
            st.text_area("", sections["ASSESSMENT"], height=200, key="a", label_visibility="collapsed")
            st.markdown("**Plan**")
            st.text_area("", sections["PLAN"], height=200, key="p", label_visibility="collapsed")

        with st.spinner("Generating billing codes..."):
            icd_prompt = f"""You are a certified medical coder with expertise in ICD-10-CM and CPT coding.

Rules:
- Always code the DIAGNOSIS, not the symptom, when a diagnosis has been established
- If no confirmed diagnosis, code the most specific symptom codes available
- Never use unspecified codes when a more specific code exists
- Do not use duplicate or overlapping codes
- Z codes for relevant history are acceptable as secondary codes
- For substance use history, always use F-category codes (F10-F19) for active or recent use disorders, or Z86.898 for remote history
- CRITICAL: Double check every code is real and accurate before returning it
- Always consider family history of cardiac disease as a relevant secondary code using Z82.49 when present
- Do not use musculoskeletal codes unless musculoskeletal cause was explicitly confirmed
- For CPT codes, base the E&M visit level on complexity of the encounter
- Only include CPT codes for procedures and tests explicitly mentioned in the plan
- For troponin lab tests, always use CPT 84484 (Troponin I, quantitative) or 84512 (Troponin T) — never 86153

Return EXACTLY 3 ICD-10-CM codes followed by EXACTLY 3 CPT codes in this format:

ICD-10-CM CODES:
Code: [code]
Name: [full official ICD-10-CM name]
Confidence: [High / Medium / Low]
Reason: [one sentence]

CPT CODES:
Code: [code]
Name: [full official CPT code name]
Confidence: [High / Medium / Low]
Reason: [one sentence]

SOAP Note:
{soap_note}"""

            response2 = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": icd_prompt}]
            )
            codes = response2.choices[0].message.content

        st.subheader("🏷 Billing Codes")
        st.text_area("", codes, height=300, label_visibility="collapsed")

        full_report = f"""AUTONOMOUS MEDICAL SCRIBE v1.0
{'='*50}

TRANSCRIPT
{'-'*50}
{transcript}

SOAP NOTE
{'-'*50}
{soap_note}

BILLING CODES
{'-'*50}
{codes}

{'='*50}
DISCLAIMER: For demonstration purposes only.
Not validated for clinical use.
Always consult a licensed physician.
"""

        st.subheader("⬇ Download Report")
        col1, col2 = st.columns(2)

        with col1:
            doc = Document()
            doc.add_heading("Autonomous Medical Scribe v1.0", 0)
            doc.add_heading("Transcript", 1)
            doc.add_paragraph(transcript)
            doc.add_heading("SOAP Note", 1)
            doc.add_paragraph(soap_note)
            doc.add_heading("Billing Codes", 1)
            doc.add_paragraph(codes)
            doc.add_paragraph("DISCLAIMER: For demonstration purposes only. Not validated for clinical use.")
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.download_button("📄 Download Word Doc", buf,
                file_name="medical_scribe_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with col2:
            pdf_buffer = io.BytesIO()
            pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter,
                rightMargin=inch, leftMargin=inch,
                topMargin=inch, bottomMargin=inch)
            styles = getSampleStyleSheet()
            story = []
            for line in full_report.split("\n"):
                if line.strip() == "":
                    story.append(Spacer(1, 6))
                else:
                    clean = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(clean, styles["Normal"]))
            pdf_doc.build(story)
            pdf_buffer.seek(0)
            st.download_button("📑 Download PDF", pdf_buffer,
                file_name="medical_scribe_report.pdf",
                mime="application/pdf")

st.markdown("---")
st.caption("⚠ For demonstration purposes only. Not validated for clinical use. Always consult a licensed physician.")
